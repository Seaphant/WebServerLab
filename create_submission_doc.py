#!/usr/bin/env python3
"""Create the Web Server Lab submission Word document."""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

LAB_DIR = os.path.dirname(os.path.abspath(__file__))

def create_lab_submission():
    doc = Document()
    
    # Title
    title = doc.add_heading('Web Server Lab - Submission', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Section 1: Lab Overview
    doc.add_heading('1. Lab Overview', level=1)
    doc.add_paragraph(
        'This lab implements a simple HTTP Web Server in Python using TCP socket programming. '
        'The server handles one HTTP request at a time, accepts and parses HTTP requests, '
        'retrieves requested files from the file system, and sends HTTP response messages to the client. '
        'If the requested file is not found, the server returns a 404 Not Found response.'
    )
    
    # Section 2: Server Code Location
    doc.add_heading('2. Server Code', level=1)
    doc.add_paragraph('The complete server code is committed to GitHub and included in this project.')
    
    # Section 3: Step-by-Step Lab Completion
    doc.add_heading('3. Lab Steps and Screenshots', level=1)
    
    # Step 1
    doc.add_heading('Step 1: Start the Web Server', level=2)
    doc.add_paragraph(
        'Run the web server with: python3 webserver.py'
    )
    terminal_img = os.path.join(LAB_DIR, 'screenshot_terminal.png')
    if os.path.exists(terminal_img):
        doc.add_picture(terminal_img, width=Inches(5.5))
    else:
        doc.add_paragraph('[INSERT SCREENSHOT 1: Terminal showing "Ready to serve..." with server running]')
    doc.add_paragraph()
    
    # Step 2
    doc.add_heading('Step 2: Access HelloWorld.html in Browser', level=2)
    doc.add_paragraph(
        'Open a web browser and navigate to: http://localhost:6789/HelloWorld.html'
    )
    doc.add_paragraph(
        '(Or use your machine IP address if testing from another host: http://YOUR_IP:6789/HelloWorld.html)'
    )
    helloworld_img = os.path.join(LAB_DIR, 'screenshot_helloworld.png')
    if os.path.exists(helloworld_img):
        doc.add_picture(helloworld_img, width=Inches(5.5))
    else:
        doc.add_paragraph('[INSERT SCREENSHOT 2: Browser displaying the HelloWorld.html content]')
    doc.add_paragraph()
    
    # Step 3
    doc.add_heading('Step 3: Verify 404 Not Found', level=2)
    doc.add_paragraph(
        'Try to access a non-existent file: http://localhost:6789/Nonexistent.html'
    )
    img_404 = os.path.join(LAB_DIR, 'screenshot_404.png')
    if os.path.exists(img_404):
        doc.add_picture(img_404, width=Inches(5.5))
    else:
        doc.add_paragraph('[INSERT SCREENSHOT 3: Browser showing "404 Not Found" message]')
    doc.add_paragraph()
    
    # Section 4: Verification
    doc.add_heading('4. Verification', level=1)
    doc.add_paragraph(
        'The screenshots above verify that:'
    )
    bullet_points = [
        'The server successfully starts and listens on port 6789',
        'The server correctly serves HTML files when requested',
        'The browser displays the full contents of HelloWorld.html',
        'The server returns proper 404 response for non-existent files'
    ]
    for point in bullet_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('--- End of Web Server Lab Submission ---')
    
    # Save
    output_path = '/Users/williamnguyen10/WebServerLab/WebServer_Lab_Submission.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_lab_submission()
